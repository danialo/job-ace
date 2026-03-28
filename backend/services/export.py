"""Resume Export Service.

Renders a ResumeDocument into PDF, DOCX, or TXT format.
Exporters are dumb — they render from the canonical model,
they do not perform semantic normalization.

See specs/resume-object-model.md for design rationale.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader
from markupsafe import Markup
from sqlalchemy.orm import Session

from backend.models import models as db_models
from backend.models.resume_document import (
    BulletsContent,
    EntryHeader,
    ItemsContent,
    ProseContent,
    ResumeDocument,
    SkillsContent,
)
from backend.services.resume_normalizer import ResumeNormalizer

TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "resume"


# ---------------------------------------------------------------------------
# HTML rendering helpers (for PDF via WeasyPrint)
# ---------------------------------------------------------------------------

def _content_to_html(content) -> str:
    """Render EntryContent as HTML."""
    if isinstance(content, BulletsContent):
        items = "".join(f"<li>{b}</li>" for b in content.bullets)
        return f"<ul>{items}</ul>"
    elif isinstance(content, ProseContent):
        return "".join(f"<p>{p}</p>" for p in content.paragraphs)
    elif isinstance(content, SkillsContent):
        parts = []
        for group in content.groups:
            if group.label:
                items_str = ", ".join(group.items)
                parts.append(f"<p><strong>{group.label}:</strong> {items_str}</p>")
            else:
                parts.append(f"<p>{', '.join(group.items)}</p>")
        return f'<div class="skills-group">{"".join(parts)}</div>'
    elif isinstance(content, ItemsContent):
        items = "".join(f"<li>{item}</li>" for item in content.items)
        return f"<ul>{items}</ul>"
    return ""


def _header_dates(header: EntryHeader | None) -> str:
    """Render dates from an EntryHeader."""
    if not header:
        return ""
    parts = []
    if header.start_date:
        parts.append(header.start_date.display())
    if header.is_current:
        parts.append("Present")
    elif header.end_date:
        parts.append(header.end_date.display())
    return " \u2013 ".join(parts)


def _content_to_txt(content) -> str:
    """Render EntryContent as plain text (ATS-safe).

    - Bullets: each prefixed with a dash
    - Prose: paragraphs separated by blank lines
    - Skills: groups as label: item1, item2 (or unlabeled inline)
    - Items: each on its own line
    """
    if isinstance(content, BulletsContent):
        return "\n".join(f"- {b}" for b in content.bullets)
    elif isinstance(content, ProseContent):
        return "\n\n".join(content.paragraphs)
    elif isinstance(content, SkillsContent):
        lines = []
        for group in content.groups:
            items_str = ", ".join(group.items)
            if group.label:
                lines.append(f"{group.label}: {items_str}")
            else:
                lines.append(items_str)
        return "\n".join(lines)
    elif isinstance(content, ItemsContent):
        return "\n".join(content.items)
    return ""


def _doc_to_template_data(doc: ResumeDocument) -> Dict:
    """Convert ResumeDocument to template-friendly dict for Jinja2."""
    contact = None
    basics = doc.basics
    if basics.name:
        contact = {
            "name": basics.name,
            "email": basics.email,
            "phone": basics.phone,
            "location": basics.location,
            "linkedin": basics.linkedin,
        }

    sections = []
    for section in doc.sections:
        blocks = []
        for entry in section.entries:
            block = {
                "job_title": entry.header.title if entry.header else None,
                "company": entry.header.organization if entry.header else None,
                "dates": _header_dates(entry.header),
                "html_content": Markup(_content_to_html(entry.content)),
            }
            blocks.append(block)
        sections.append({"heading": section.heading, "blocks": blocks})

    return {"contact": contact, "sections": sections}


# ---------------------------------------------------------------------------
# ExportService
# ---------------------------------------------------------------------------

class ExportService:
    def __init__(self, db: Session):
        self.db = db

    def list_templates(self) -> List[Dict]:
        """Scan template directory for available templates."""
        templates = []
        if not TEMPLATE_DIR.is_dir():
            return templates
        for meta_path in TEMPLATE_DIR.glob("*/meta.json"):
            try:
                meta = json.loads(meta_path.read_text(encoding="utf-8"))
                meta["id"] = meta_path.parent.name
                templates.append(meta)
            except (json.JSONDecodeError, OSError):
                continue
        return templates

    def render_pdf(
        self,
        job_id: int,
        block_ids: List[int],
        template: str = "classic",
        version: str = "v1",
    ) -> bytes:
        """Render resume blocks as a PDF via HTML+WeasyPrint."""
        import weasyprint

        doc = self._build_document(block_ids)
        render_data = _doc_to_template_data(doc)

        template_dir = TEMPLATE_DIR / template
        if not template_dir.is_dir():
            raise ValueError(f"Template '{template}' not found")

        env = Environment(loader=FileSystemLoader(str(template_dir)), autoescape=True)
        tmpl = env.get_template("template.html")
        html_string = tmpl.render(**render_data)

        css_path = template_dir / "style.css"
        stylesheets = [weasyprint.CSS(filename=str(css_path))] if css_path.exists() else []

        html_doc = weasyprint.HTML(string=html_string, base_url=str(template_dir))
        return html_doc.write_pdf(stylesheets=stylesheets)

    def render_docx(
        self,
        job_id: int,
        block_ids: List[int],
        template: str = "classic",
        version: str = "v1",
    ) -> bytes:
        """Build a DOCX resume programmatically with python-docx."""
        resume = self._build_document(block_ids)
        doc = Document()

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)

        # Contact header
        basics = resume.basics
        if basics.name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(basics.name)
            run.bold = True
            run.font.size = Pt(16)

        info_parts = [v for v in (basics.email, basics.phone, basics.location, basics.linkedin) if v]
        if info_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(info_parts))
            run.font.size = Pt(9.5)

        # Sections
        for section in resume.sections:
            heading_p = doc.add_heading(section.heading, level=2)
            heading_p.style.font.size = Pt(11)

            for entry in section.entries:
                # Entry header
                if entry.header and (entry.header.title or entry.header.organization):
                    p = doc.add_paragraph()
                    if entry.header.title:
                        run = p.add_run(entry.header.title)
                        run.bold = True
                    if entry.header.organization:
                        run = p.add_run(f" — {entry.header.organization}")
                        run.italic = True
                    dates = _header_dates(entry.header)
                    if dates:
                        run = p.add_run(f"    {dates}")
                        run.font.size = Pt(9.5)

                # Content — render by type
                content = entry.content
                if isinstance(content, BulletsContent):
                    for bullet in content.bullets:
                        doc.add_paragraph(bullet, style="List Bullet")
                elif isinstance(content, ProseContent):
                    for para in content.paragraphs:
                        doc.add_paragraph(para)
                elif isinstance(content, SkillsContent):
                    for group in content.groups:
                        text = ", ".join(group.items)
                        if group.label:
                            p = doc.add_paragraph()
                            run = p.add_run(f"{group.label}: ")
                            run.bold = True
                            p.add_run(text)
                        else:
                            doc.add_paragraph(text)
                elif isinstance(content, ItemsContent):
                    for item in content.items:
                        doc.add_paragraph(item, style="List Bullet")

        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_txt(
        self,
        job_id: int,
        block_ids: List[int],
        version: str = "v1",
    ) -> str:
        """Render resume blocks as plain text (ATS-safe).

        Output is human-readable and machine-parseable:
        - Name and contact info as pipe-separated line
        - Section headings in ALL CAPS with underline dashes
        - Entry headers as title/org (dates)
        - Bullets prefixed with dash
        - Skills as label: item1, item2
        - Items one per line
        """
        resume = self._build_document(block_ids, job_id=job_id)
        lines: List[str] = []

        # Contact header
        basics = resume.basics
        if basics.name:
            lines.append(basics.name)
        info_parts = [v for v in (basics.email, basics.phone, basics.location, basics.linkedin) if v]
        if info_parts:
            lines.append(" | ".join(info_parts))
        if lines:
            lines.append("")  # blank line after header

        # Sections
        for section in resume.sections:
            heading = section.heading.upper()
            lines.append(heading)
            lines.append("-" * len(heading))

            for entry in section.entries:
                # Entry header line
                if entry.header and (entry.header.title or entry.header.organization):
                    header_parts = []
                    if entry.header.title:
                        header_parts.append(entry.header.title)
                    if entry.header.organization:
                        header_parts.append(entry.header.organization)
                    header_str = " | ".join(header_parts)
                    dates = _header_dates(entry.header)
                    if dates:
                        header_str = f"{header_str} ({dates})"
                    lines.append(header_str)

                # Content
                content_txt = _content_to_txt(entry.content)
                if content_txt:
                    lines.append(content_txt)
                lines.append("")  # blank line between entries

            lines.append("")  # extra blank line between sections

        # Strip trailing blank lines, add single trailing newline
        result = "\n".join(lines).rstrip() + "\n"
        return result

    def _build_document(
        self,
        block_ids: List[int],
        job_id: Optional[int] = None,
    ) -> ResumeDocument:
        """Normalize blocks into a ResumeDocument.

        If job_id is provided, persists the canonical document as
        derived/resume_document.json for debuggability and reproducibility.
        """
        normalizer = ResumeNormalizer(self.db)
        doc = normalizer.normalize(block_ids)

        if job_id is not None:
            self._persist_document(job_id, doc)

        return doc

    def _persist_document(self, job_id: int, doc: ResumeDocument) -> None:
        """Persist the canonical ResumeDocument as a JSON artifact."""
        from backend.services.artifacts import ArtifactManager

        job_posting = self.db.get(db_models.JobPosting, job_id)
        if job_posting is None:
            return  # no job context — skip persistence silently

        artifact_mgr = ArtifactManager(self.db)
        artifact_mgr.write_text(
            job_posting,
            "resume_document",
            "derived/resume_document.json",
            doc.model_dump_json(indent=2),
        )
